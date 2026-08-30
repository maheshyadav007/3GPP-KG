from __future__ import annotations

import io
import subprocess
import zipfile
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from threegpp_kg.config import load_settings
from threegpp_kg.constants import BlockKind, Conclusion
from threegpp_kg.ingestion.chunking import build_chunks
from threegpp_kg.ingestion.normalize import normalize_organization
from threegpp_kg.parsers import (
    documents,
    parse_docx,
    parse_tdoc_workbook,
    parse_tdoc_workbook_package,
)
from threegpp_kg.parsers.documents import (
    NoExtractableTextError,
    UnsafeDocumentError,
    UnsupportedDocumentError,
    parse_document,
)


def workbook_bytes(tdoc_id: str = "R2-2600001") -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "TDoc_List"
    sheet.append(
        [
            "TDoc",
            "Title",
            "Source",
            "Type",
            "For",
            "TDoc Status",
            "Is revision of",
            "Revised to",
            "Release",
            "Spec",
            "Related WIs",
        ]
    )
    sheet.append(
        [
            tdoc_id,
            "Paging relaxation",
            "Qualcomm, Ericsson",
            "discussion",
            "Agreement",
            "revised",
            None,
            "R2-2601001",
            "Rel-20",
            "38.331",
            "NR_Core",
        ]
    )
    private = workbook.create_sheet("Registration")
    private.append(["email", "person@example.com"])
    stream = io.BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("1 Introduction", level=1)
    document.add_paragraph("Discussion: Companies considered two alternatives.")
    document.add_paragraph("Agreement: Adopt the revised text.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Spec"
    table.rows[0].cells[1].text = "38.331"
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def strict_docx_bytes() -> bytes:
    replacements = {
        transitional: strict
        for strict, transitional in documents._STRICT_OOXML_NAMESPACES.items()
    }
    output = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(docx_bytes())) as source,
        zipfile.ZipFile(output, "w") as target,
    ):
        for member in source.infolist():
            payload = source.read(member)
            if member.filename.lower().endswith((".xml", ".rels")):
                for transitional, strict in replacements.items():
                    payload = payload.replace(transitional, strict)
            target.writestr(member, payload)
    return output.getvalue()


def pptx_bytes() -> bytes:
    stream = io.BytesIO()
    slide = """<?xml version="1.0" encoding="UTF-8"?>
    <p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
      xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
      <p:cSld><p:spTree>
        <p:sp><p:nvSpPr><p:cNvPr id="1" name="Title"/><p:cNvSpPr/><p:nvPr>
          <p:ph type="title"/></p:nvPr></p:nvSpPr><p:txBody>
          <a:p><a:r><a:t>Multiple reference configuration</a:t></a:r></a:p>
        </p:txBody></p:sp>
        <p:sp><p:nvSpPr><p:cNvPr id="2" name="Body"/><p:cNvSpPr/><p:nvPr/>
          </p:nvSpPr><p:txBody><a:p><a:pPr lvl="0"/><a:r>
          <a:t>Agreement: adopt option A.</a:t></a:r></a:p></p:txBody></p:sp>
        <p:graphicFrame><a:graphic><a:graphicData><a:tbl><a:tr>
          <a:tc><a:txBody><a:p><a:r><a:t>Spec</a:t></a:r></a:p></a:txBody></a:tc>
          <a:tc><a:txBody><a:p><a:r><a:t>38.331</a:t></a:r></a:p></a:txBody></a:tc>
        </a:tr></a:tbl></a:graphicData></a:graphic></p:graphicFrame>
      </p:spTree></p:cSld>
    </p:sld>"""
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("ppt/slides/slide1.xml", slide)
    return stream.getvalue()


def test_workbook_parser_normalizes_and_excludes_private_sheets() -> None:
    parsed = parse_tdoc_workbook(workbook_bytes(), "RAN2-133", "https://www.3gpp.org/list.xlsx")
    assert len(parsed) == 1
    tdoc = parsed[0].tdocs[0]
    assert tdoc.id == "R2-2600001"
    assert tdoc.status == Conclusion.REVISED
    assert tdoc.revised_to == "R2-2601001"
    assert tdoc.specifications == ["38.331"]


def test_ct1_workbook_identifier_is_recognized() -> None:
    parsed = parse_tdoc_workbook(
        workbook_bytes("C1-261234"),
        "CT1-162",
        "https://www.3gpp.org/C1-index.xlsx",
    )
    assert parsed[0].tdocs[0].id == "C1-261234"


def test_tdoc_workbook_package_extracts_xlsx_from_zip() -> None:
    package = io.BytesIO()
    with zipfile.ZipFile(package, "w") as archive:
        archive.writestr("SA2-173_Index_2026.xlsx", workbook_bytes("S2-261234"))
    parsed = parse_tdoc_workbook_package(
        package.getvalue(),
        "SA2-173_Index_2026.zip",
        "SA2-173",
        "https://www.3gpp.org/SA2-173_Index_2026.zip",
    )
    assert parsed[0].tdocs[0].id == "S2-261234"


def test_tdoc_workbook_package_rejects_path_traversal() -> None:
    package = io.BytesIO()
    with zipfile.ZipFile(package, "w") as archive:
        archive.writestr("../Index.xlsx", workbook_bytes())
    with pytest.raises(UnsafeDocumentError, match="unsafe TDoc index archive member"):
        parse_tdoc_workbook_package(
            package.getvalue(),
            "Index.zip",
            "RAN2-133",
            "https://www.3gpp.org/Index.zip",
        )


def test_docx_parser_preserves_semantic_blocks_and_chunks() -> None:
    blocks = parse_docx(docx_bytes(), "R2-2600001")
    assert [block.kind for block in blocks[:3]] == [
        BlockKind.HEADING,
        BlockKind.DISCUSSION,
        BlockKind.AGREEMENT,
    ]
    assert any(block.kind == BlockKind.TABLE_ROW for block in blocks)
    chunks = build_chunks(
        blocks,
        load_settings().chunking.model_copy(
            update={"min_tokens": 1, "target_tokens": 5, "max_tokens": 30}
        ),
    )
    assert chunks and all(chunk.block_ids for chunk in chunks)
    assert len({chunk.id for chunk in chunks}) == len(chunks)


def test_strict_ooxml_docx_is_normalized_before_parsing() -> None:
    blocks = parse_document(strict_docx_bytes(), "R2-1.docx", "R2-1")

    assert blocks
    assert blocks[0].text == "1 Introduction"


def test_text_marker_is_preserved_as_evidence() -> None:
    blocks = parse_document(
        b"S2-1 withdrawn and replaced by S2-2\n",
        "S2-1.txt",
        "S2-1",
    )

    assert [block.text for block in blocks] == ["S2-1 withdrawn and replaced by S2-2"]


def test_organization_alias_normalization_is_explicit() -> None:
    aliases = {"qualcomm incorporated": "Qualcomm"}
    assert normalize_organization(" Qualcomm Incorporated ", aliases) == "Qualcomm"
    assert normalize_organization("New Company", aliases) == "New Company"


def test_tdoc_zip_extracts_supported_document_and_reindexes_blocks() -> None:
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("nested/R2-2600001.docx", docx_bytes())
        archive.writestr("__MACOSX/._R2-2600001.docx", b"resource fork")
        archive.writestr("readme.txt", "ignored")
    blocks = parse_document(stream.getvalue(), "R2-2600001.zip", "R2-2600001")
    assert blocks[0].document_id == "R2-2600001"
    assert [block.index for block in blocks] == list(range(len(blocks)))


@pytest.mark.parametrize(
    "attachment_name", ["TDoc_List_Meeting_RAN2#131.xlsx", "PartList_RAN2#131.xlsx"]
)
def test_document_archive_skips_bundled_meeting_export_workbooks(
    attachment_name: str,
) -> None:
    workbook = Workbook()
    workbook.active.append(["first"])
    workbook.active.append(["second"])
    workbook_stream = io.BytesIO()
    workbook.save(workbook_stream)
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("R2-2600001.docx", docx_bytes())
        archive.writestr(attachment_name, workbook_stream.getvalue())

    blocks = parse_document(
        stream.getvalue(),
        "R2-2600001.zip",
        "R2-2600001",
        load_settings().parsers.model_copy(update={"max_workbook_rows": 1}),
    )

    assert blocks
    assert all(block.section_path != ["Sheet"] for block in blocks)


def test_pptx_parser_preserves_slides_lists_and_tables() -> None:
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("R2-2409031.pptx", pptx_bytes())
    blocks = parse_document(stream.getvalue(), "R2-2409031.zip", "R2-2409031")
    assert [block.kind for block in blocks] == [
        BlockKind.HEADING,
        BlockKind.LIST_ITEM,
        BlockKind.TABLE_ROW,
    ]
    assert blocks[1].section_path == ["Slide 1", "Multiple reference configuration"]
    assert blocks[2].text == "Spec | 38.331"


def test_xlsx_parser_preserves_rows_and_excludes_private_or_hidden_sheets() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Resolutions"
    sheet.append(["Issue", "Status", "CR"])
    sheet.append([17, "Agreed", "R2-2507008"])
    private = workbook.create_sheet("Registration")
    private.append(["person@example.com"])
    hidden = workbook.create_sheet("Internal")
    hidden.sheet_state = "hidden"
    hidden.append(["not public evidence"])
    stream = io.BytesIO()
    workbook.save(stream)

    blocks = parse_document(stream.getvalue(), "R2-2507008.xlsx", "R2-2507008")

    assert [block.text for block in blocks] == [
        "Issue | Status | CR",
        "17 | Agreed | R2-2507008",
    ]
    assert all(block.kind == BlockKind.TABLE_ROW for block in blocks)
    assert all(block.section_path == ["Resolutions"] for block in blocks)
    assert [block.table_row for block in blocks] == [1, 2]


def test_document_archive_parses_xlsx_and_one_nested_zip() -> None:
    workbook = Workbook()
    workbook.active.append(["Issue", "Resolution"])
    workbook.active.append([1, "Adopt option A"])
    workbook_stream = io.BytesIO()
    workbook.save(workbook_stream)

    nested = io.BytesIO()
    with zipfile.ZipFile(nested, "w") as archive:
        archive.writestr("review.xlsx", workbook_stream.getvalue())
    outer = io.BytesIO()
    with zipfile.ZipFile(outer, "w") as archive:
        archive.writestr("review.zip", nested.getvalue())

    blocks = parse_document(outer.getvalue(), "R2-1.zip", "R2-1")

    assert [block.text for block in blocks] == ["Issue | Resolution", "1 | Adopt option A"]
    assert [block.index for block in blocks] == [0, 1]


def test_document_archive_rejects_excessive_nesting() -> None:
    payload = io.BytesIO()
    with zipfile.ZipFile(payload, "w") as archive:
        archive.writestr("document.docx", docx_bytes())
    for name in ("inner.zip", "middle.zip", "outer.zip"):
        wrapped = io.BytesIO()
        with zipfile.ZipFile(wrapped, "w") as archive:
            archive.writestr(name, payload.getvalue())
        payload = wrapped

    with pytest.raises(UnsafeDocumentError, match="safe depth"):
        parse_document(payload.getvalue(), "R2-1.zip", "R2-1")


def test_document_parser_enforces_configured_workbook_limits() -> None:
    workbook = Workbook()
    workbook.active.append(["Header", None, "Third column"])
    workbook.active.append(["Second row"])
    stream = io.BytesIO()
    workbook.save(stream)
    content = stream.getvalue()
    defaults = load_settings().parsers

    with pytest.raises(UnsafeDocumentError, match="row limit"):
        parse_document(
            content,
            "R2-1.xlsx",
            "R2-1",
            defaults.model_copy(update={"max_workbook_rows": 1}),
        )
    with pytest.raises(UnsafeDocumentError, match="block limit"):
        parse_document(
            content,
            "R2-1.xlsx",
            "R2-1",
            defaults.model_copy(update={"max_document_blocks": 1}),
        )


def test_document_parser_enforces_configured_archive_limits() -> None:
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("first.docx", docx_bytes())
        archive.writestr("second.docx", docx_bytes())
    defaults = load_settings().parsers

    with pytest.raises(UnsafeDocumentError, match="too many members"):
        parse_document(
            stream.getvalue(),
            "R2-1.zip",
            "R2-1",
            defaults.model_copy(update={"max_archive_members": 1}),
        )
    with pytest.raises(UnsafeDocumentError, match="expands beyond"):
        parse_document(
            stream.getvalue(),
            "R2-1.zip",
            "R2-1",
            defaults.model_copy(update={"max_archive_uncompressed_bytes": 1024}),
        )


@pytest.mark.parametrize("filename", ["legacy.doc", "unknown.bin", "no-extension"])
def test_unsupported_document_formats_fail_closed(filename: str) -> None:
    with pytest.raises(UnsupportedDocumentError):
        parse_document(b"not a document", filename, "R2-1")


def test_document_archive_rejects_traversal_invalid_and_empty_content() -> None:
    traversal = io.BytesIO()
    with zipfile.ZipFile(traversal, "w") as archive:
        archive.writestr("../escape.docx", docx_bytes())
    with pytest.raises(UnsafeDocumentError, match="unsafe"):
        parse_document(traversal.getvalue(), "R2-1.zip", "R2-1")
    with pytest.raises(UnsafeDocumentError, match="invalid"):
        parse_document(b"not-a-zip", "R2-1.zip", "R2-1")

    empty = io.BytesIO()
    with zipfile.ZipFile(empty, "w") as archive:
        archive.writestr("readme.bin", "nothing supported")
    with pytest.raises(UnsupportedDocumentError, match="no supported"):
        parse_document(empty.getvalue(), "R2-1.zip", "R2-1")

    image_only = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    pdf = io.BytesIO()
    writer.write(pdf)
    with zipfile.ZipFile(image_only, "w") as archive:
        archive.writestr("scan.pdf", pdf.getvalue())
    with pytest.raises(NoExtractableTextError, match="no extractable text"):
        parse_document(image_only.getvalue(), "R2-3.zip", "R2-3")

    no_members = io.BytesIO()
    with zipfile.ZipFile(no_members, "w"):
        pass
    with pytest.raises(UnsafeDocumentError, match="empty.*quarantined"):
        parse_document(no_members.getvalue(), "R2-2.zip", "R2-2")


def test_document_archive_reports_legacy_only_content() -> None:
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("R2-1.doc", b"legacy")
    with pytest.raises(UnsupportedDocumentError, match="legacy DOC"):
        parse_document(stream.getvalue(), "R2-1.zip", "R2-1")


def test_legacy_doc_uses_configured_converter(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("R2-1.doc", b"legacy")

    monkeypatch.setattr(documents.shutil, "which", lambda name: f"/usr/bin/{name}")

    def convert(command: list[str], **_kwargs: object) -> subprocess.CompletedProcess[bytes]:
        output = Path(command[command.index("-output") + 1])
        output.write_bytes(docx_bytes())
        return subprocess.CompletedProcess(command, 0, b"", b"")

    monkeypatch.setattr(documents.subprocess, "run", convert)
    config = load_settings().parsers.model_copy(update={"legacy_converter": "textutil"})
    blocks = parse_document(stream.getvalue(), "R2-1.zip", "R2-1", config)
    assert blocks and blocks[0].document_id == "R2-1"


def test_pdf_dispatch_accepts_valid_empty_pdf() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    stream = io.BytesIO()
    writer.write(stream)
    assert parse_document(stream.getvalue(), "R2-1.pdf", "R2-1") == []


def test_oversized_blocks_are_split_into_stable_bounded_chunks() -> None:
    blocks = parse_docx(docx_bytes(), "R2-1")
    oversized = blocks[1].model_copy(update={"text": "word " * 100})
    chunks = build_chunks(
        [oversized],
        load_settings().chunking.model_copy(
            update={"min_tokens": 2, "target_tokens": 4, "max_tokens": 10}
        ),
    )
    assert len(chunks) > 1
    assert all(chunk.token_count <= 10 for chunk in chunks)
