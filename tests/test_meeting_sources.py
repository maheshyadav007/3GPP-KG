from __future__ import annotations

import hashlib
import io
from datetime import UTC, date, datetime, timedelta

import pytest
from docx import Document
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from threegpp_kg.config import load_settings
from threegpp_kg.constants import (
    BlockKind,
    Conclusion,
    DocumentState,
    EvidenceAuthority,
    ObservationType,
    SourceRole,
)
from threegpp_kg.domain import (
    DocumentBlock,
    EvidenceRef,
    Meeting,
    MeetingObservation,
    MeetingSource,
    TDoc,
)
from threegpp_kg.ingestion.download import DownloadedArtifact
from threegpp_kg.meeting_sources import extract_meeting_observations, ingest_meeting_source
from threegpp_kg.repository import InMemoryRepository, SqlRepository
from threegpp_kg.service import KnowledgeService
from threegpp_kg.storage.database import (
    Base,
    KnowledgeEdgeRow,
    MeetingObservationRow,
)
from threegpp_kg.storage.object_store import LocalObjectStore


def _chair_note() -> bytes:
    document = Document()
    document.add_heading("7 Mobility", level=1)
    document.add_paragraph("Discussion: R2-2600001 requires one further revision.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Deadline:"
    table.cell(0, 1).text = "Provide the revision at the next meeting."
    document.add_paragraph("Agreement: R2-2600001 was agreed for TS 38.331.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_post_meeting_rollup_splits_structured_discussions() -> None:
    block = DocumentBlock(
        id="block-post",
        document_id="post:RAN2-132",
        index=0,
        kind=BlockKind.PARAGRAPH,
        text=(
            "[POST132][001] Intended outcome: agree R2-2509001. Deadline: Short "
            "[POST132][002] Intended outcome: revise R2-2509002. Deadline: Long"
        ),
    )
    observations = extract_meeting_observations(
        [block],
        meeting_id="RAN2-132",
        artifact_version_id="artifact-post",
        source_role=SourceRole.POST_MEETING_DISCUSSION,
        authority=EvidenceAuthority.POST_MEETING_DISCUSSION,
        evidence_by_block={block.id: "ev-post"},
    )
    assert len(observations) == 4
    assert {item.tdoc_ids[0] for item in observations} == {
        "R2-2509001",
        "R2-2509002",
    }


@pytest.mark.asyncio
async def test_ingested_chair_note_persists_observations_graph_and_briefing(tmp_path) -> None:
    engine = create_async_engine("sqlite+aiosqlite://")
    sessions = async_sessionmaker(engine, expire_on_commit=False)
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)
    meeting = Meeting(
        id="RAN2-132",
        working_group_id="RAN2",
        number=132,
        name="RAN2 #132",
        source_url="https://www.3gpp.org/meeting",
        starts_on=date(2026, 2, 16),
        ends_on=date(2026, 2, 20),
    )
    content = _chair_note()
    artifact = DownloadedArtifact(
        url="https://www.3gpp.org/meeting/Inbox/Chair_Notes/final.docx",
        content=content,
        sha256=hashlib.sha256(content).hexdigest(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        etag=None,
        last_modified="Fri, 20 Feb 2026 17:00:00 GMT",
    )
    settings = load_settings()
    async with sessions() as session:
        blocks, chunks, observations = await ingest_meeting_source(
            session,
            LocalObjectStore(tmp_path / "objects"),
            "source-fixture",
            meeting,
            "R2_132_ChairNotes_final_clean.docx",
            artifact,
            SourceRole.CHAIR_NOTES,
            settings.chunking,
            evidence_blocks=settings.evidence_blocks,
            parser_config=settings.parsers,
        )
        await session.commit()
        observation_count = await session.scalar(select(func.count(MeetingObservationRow.id)))
        observation_edges = await session.scalar(
            select(func.count(KnowledgeEdgeRow.id)).where(
                KnowledgeEdgeRow.source_type == "observation"
            )
        )
    assert len(blocks) == 3
    assert "Deadline:" in blocks[1].text
    assert blocks[2].text.startswith("Agreement:")
    assert chunks
    assert observation_count == len(observations) >= 3
    assert observation_edges
    assert {item.observation_type for item in observations} >= {
        ObservationType.DECISION,
        ObservationType.DISCUSSION_SUMMARY,
        ObservationType.FOLLOW_UP_ACTION,
    }

    service = KnowledgeService(SqlRepository(sessions, preview_dataset_version="source-fixture"))
    briefing = await service.meeting_briefing(meeting.id)
    assert briefing.data is not None
    assert briefing.data.decisions[0].tdoc_ids == ["R2-2600001"]
    assert briefing.data.sources[0].document_state == DocumentState.FINAL_CLEAN
    assert briefing.evidence
    await engine.dispose()


@pytest.mark.asyncio
async def test_briefing_diffs_versions_and_newsletter_consumes_observations() -> None:
    meeting = Meeting(
        id="RAN2-132",
        working_group_id="RAN2",
        number=132,
        name="RAN2 #132",
        source_url="https://www.3gpp.org/meeting",
        starts_on=date(2026, 2, 16),
        ends_on=date(2026, 2, 20),
        readiness="final_ready",
    )
    observed = datetime(2026, 2, 20, tzinfo=UTC)
    sources = [
        MeetingSource(
            artifact_version_id=f"artifact-{version}",
            meeting_id=meeting.id,
            source_role=SourceRole.CHAIR_NOTES,
            logical_document_id=f"chair_notes:{meeting.id}",
            document_id=f"chair_notes:{meeting.id}:{version}",
            filename=f"chair-{version}.docx",
            source_url=f"https://example.test/chair-{version}.docx",
            sha256=version * 64,
            document_state=DocumentState.FINAL_CLEAN,
            authority=EvidenceAuthority.CHAIR_NOTES,
            observed_at=observed + timedelta(hours=version == "b"),
        )
        for version in ("a", "b")
    ]
    before = MeetingObservation(
        id="observation-before",
        meeting_id=meeting.id,
        artifact_version_id="artifact-a",
        source_role=SourceRole.CHAIR_NOTES,
        authority=EvidenceAuthority.CHAIR_NOTES,
        observation_type=ObservationType.DECISION,
        observation_key="same-decision",
        text="Agreement: R2-2600001 was postponed.",
        tdoc_ids=["R2-2600001"],
        conclusion=Conclusion.POSTPONED,
        evidence_ids=["ev-old"],
        content_hash="old",
    )
    after = before.model_copy(
        update={
            "id": "observation-after",
            "artifact_version_id": "artifact-b",
            "text": "Agreement: R2-2600001 was agreed.",
            "conclusion": Conclusion.AGREED,
            "evidence_ids": ["ev-new"],
            "content_hash": "new",
        }
    )
    evidence = EvidenceRef(
        id="ev-new",
        source_url=sources[-1].source_url,
        artifact_sha256=sources[-1].sha256,
        authority=EvidenceAuthority.CHAIR_NOTES,
        meeting_id=meeting.id,
        excerpt=after.text,
    )
    tdoc = TDoc(
        id="R2-2600001",
        meeting_id=meeting.id,
        title="Mobility update",
        status=Conclusion.AGREED,
        evidence_ids=["ev-new"],
    )
    service = KnowledgeService(
        InMemoryRepository(
            [meeting],
            [tdoc],
            [evidence],
            dataset_version="diff-fixture",
            sources=sources,
            observations=[before, after],
        )
    )
    briefing = await service.meeting_briefing(meeting.id)
    assert briefing.data is not None
    assert briefing.data.observations == [after]
    assert briefing.data.changes[0].change_type == "changed"
    packet = await service.newsletter_packet(meeting.id)
    assert packet.data is not None
    assert any(
        signal.category == "meeting_decision" and signal.evidence_ids == ["ev-new"]
        for signal in packet.data.signals
    )


@pytest.mark.asyncio
async def test_final_report_supersedes_matching_chair_observation() -> None:
    meeting = Meeting(
        id="RAN2-132",
        working_group_id="RAN2",
        number=132,
        name="RAN2 #132",
        source_url="https://www.3gpp.org/meeting",
    )
    now = datetime(2026, 2, 20, tzinfo=UTC)
    sources = [
        MeetingSource(
            artifact_version_id=f"artifact-{role}",
            meeting_id=meeting.id,
            source_role=role,
            logical_document_id=f"{role}:{meeting.id}",
            document_id=f"{role}:{meeting.id}:v1",
            filename=f"{role}.docx",
            source_url=f"https://example.test/{role}.docx",
            sha256=("a" if role == SourceRole.CHAIR_NOTES else "b") * 64,
            authority=authority,
            observed_at=now,
        )
        for role, authority in (
            (SourceRole.CHAIR_NOTES, EvidenceAuthority.CHAIR_NOTES),
            (SourceRole.REPORT, EvidenceAuthority.FINAL_REPORT),
        )
    ]
    observations = [
        MeetingObservation(
            id=f"observation-{role}",
            meeting_id=meeting.id,
            artifact_version_id=f"artifact-{role}",
            source_role=role,
            authority=authority,
            observation_type=ObservationType.DECISION,
            observation_key="decision-key",
            text="Approved",
            tdoc_ids=["R2-2509001"],
            conclusion=Conclusion.APPROVED,
            evidence_ids=[evidence_id],
            content_hash="same-content",
        )
        for role, authority, evidence_id in (
            (SourceRole.CHAIR_NOTES, EvidenceAuthority.CHAIR_NOTES, "ev-chair"),
            (SourceRole.REPORT, EvidenceAuthority.FINAL_REPORT, "ev-report"),
        )
    ]
    service = KnowledgeService(
        InMemoryRepository(
            [meeting],
            sources=sources,
            observations=observations,
        )
    )
    provisional = await service.meeting_briefing(meeting.id, "provisional")
    final = await service.meeting_briefing(meeting.id, "final")
    assert provisional.data and provisional.data.observations == [observations[0]]
    assert final.data and final.data.observations == [observations[1]]
