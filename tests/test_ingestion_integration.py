from __future__ import annotations

import hashlib
import io
from datetime import date

import pytest
from docx import Document
from openpyxl import Workbook
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from threegpp_kg.config import ChunkingConfig
from threegpp_kg.constants import DatasetState
from threegpp_kg.domain import Meeting, SearchFilters, SearchRequest, TemporalScope
from threegpp_kg.ingestion.download import DownloadedArtifact
from threegpp_kg.ingestion.pipeline import (
    ingest_document_artifact,
    ingest_tdoc_workbook,
    update_meeting_dates,
    validate_dataset,
)
from threegpp_kg.publisher import activate_dataset
from threegpp_kg.repository import SqlRepository
from threegpp_kg.service import KnowledgeService
from threegpp_kg.storage.database import (
    ArtifactVersionRow,
    Base,
    DatasetVersionRow,
    DocumentBlockRow,
    EvidenceRow,
    KnowledgeEdgeRow,
    KnowledgeNodeRow,
    MeetingRow,
    RetrievalChunkRow,
    TDocRow,
)
from threegpp_kg.storage.object_store import LocalObjectStore


def workbook_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["TDoc", "Title", "Source", "Result", "Spec", "Release"])
    sheet.append(["R2-2600001", "Carrier aggregation", "Qualcomm", "Agreed", "38.306", "Rel-20"])
    sheet.append(["R2-2600002", "Sidelink", "Ericsson", "Revised", "38.331", "Rel-20"])
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Carrier aggregation", level=1)
    document.add_paragraph("Discussion: The companies reviewed simultaneous transmission.")
    document.add_paragraph("Conclusion: The proposal was agreed.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def artifact(url: str, content: bytes, content_type: str) -> DownloadedArtifact:
    return DownloadedArtifact(
        url=url,
        content=content,
        sha256=hashlib.sha256(content).hexdigest(),
        content_type=content_type,
        etag='"fixture"',
        last_modified="Sat, 29 Aug 2026 00:00:00 GMT",
    )


@pytest.mark.asyncio
async def test_same_source_artifact_can_belong_to_multiple_dataset_versions(tmp_path) -> None:
    engine = create_async_engine("sqlite+aiosqlite://")
    sessions = async_sessionmaker(engine, expire_on_commit=False)
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)
    meeting = Meeting(
        id="RAN2-134",
        working_group_id="RAN2",
        number=134,
        name="RAN2 #134",
        source_url="https://www.3gpp.org/meeting",
        starts_on=date(2026, 5, 18),
        ends_on=date(2026, 5, 22),
        readiness="final_ready",
    )
    store = LocalObjectStore(tmp_path / "objects")
    workbook = artifact(
        "https://www.3gpp.org/TDoc_List.xlsx",
        workbook_bytes(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    async with sessions() as session:
        await ingest_tdoc_workbook(
            session, store, "snapshot-1", meeting, "TDoc_List.xlsx", workbook
        )
        await ingest_tdoc_workbook(
            session, store, "snapshot-2", meeting, "TDoc_List.xlsx", workbook
        )
        await session.commit()

    async with sessions() as session:
        artifacts = (
            await session.scalars(
                select(ArtifactVersionRow).order_by(ArtifactVersionRow.dataset_version_id)
            )
        ).all()
        assert [row.dataset_version_id for row in artifacts] == ["snapshot-1", "snapshot-2"]
        assert artifacts[0].source_url == artifacts[1].source_url
        assert artifacts[0].sha256 == artifacts[1].sha256
        assert await session.scalar(select(func.count(TDocRow.row_id))) == 4
        assert await session.scalar(select(func.count(EvidenceRow.id))) == 4
    await engine.dispose()


@pytest.mark.asyncio
async def test_meeting_date_fallback_propagates_to_evidence_and_graph(tmp_path) -> None:
    engine = create_async_engine("sqlite+aiosqlite://")
    sessions = async_sessionmaker(engine, expire_on_commit=False)
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)
    meeting = Meeting(
        id="RAN2-127-bis",
        working_group_id="RAN2",
        number=127,
        variant="bis",
        name="RAN2 #127bis",
        source_url="https://www.3gpp.org/meeting",
        readiness="final_ready",
    )
    store = LocalObjectStore(tmp_path / "objects")
    workbook = artifact(
        "https://www.3gpp.org/TDoc_List.xlsx",
        workbook_bytes(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    async with sessions() as session:
        await ingest_tdoc_workbook(session, store, "snapshot", meeting, "TDoc_List.xlsx", workbook)
        await update_meeting_dates(
            session,
            "snapshot",
            meeting.id,
            date(2024, 10, 14),
            date(2024, 10, 18),
        )
        await session.commit()

    async with sessions() as session:
        stored_meeting = await session.scalar(select(MeetingRow))
        assert stored_meeting and stored_meeting.starts_on == date(2024, 10, 14)
        assert stored_meeting.ends_on == date(2024, 10, 18)
        evidence_dates = set(await session.scalars(select(EvidenceRow.meeting_time)))
        edge_dates = set(await session.scalars(select(KnowledgeEdgeRow.valid_at)))
        assert evidence_dates == {date(2024, 10, 18)}
        assert edge_dates == {date(2024, 10, 18)}
    await engine.dispose()


@pytest.mark.asyncio
async def test_idempotent_ingestion_document_blocks_and_atomic_activation(tmp_path) -> None:
    engine = create_async_engine("sqlite+aiosqlite://")
    sessions = async_sessionmaker(engine, expire_on_commit=False)
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)
    meeting = Meeting(
        id="RAN2-133",
        working_group_id="RAN2",
        number=133,
        name="RAN2 #133",
        source_url="https://www.3gpp.org/meeting",
        starts_on=date(2026, 5, 4),
        ends_on=date(2026, 5, 8),
        readiness="final_ready",
    )
    store = LocalObjectStore(tmp_path / "objects")
    workbook = artifact(
        "https://www.3gpp.org/TDoc_List.xlsx",
        workbook_bytes(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    document = artifact(
        "https://www.3gpp.org/R2-2600001.docx",
        docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    async with sessions() as session:
        for attempt in range(2):
            await ingest_tdoc_workbook(session, store, "v1", meeting, "TDoc_List.xlsx", workbook)
            await ingest_document_artifact(
                session,
                store,
                "v1",
                meeting,
                "R2-2600001",
                "R2-2600001.docx",
                document,
                ChunkingConfig(min_tokens=2, target_tokens=8, max_tokens=20),
            )
            if attempt == 0:
                await session.flush()
                artifact_row = await session.scalar(
                    select(ArtifactVersionRow).where(ArtifactVersionRow.kind == "tdoc")
                )
                assert artifact_row
                artifact_row.parse_status = "failed"
                artifact_row.parse_error = "simulated old parser failure"
        stats = await validate_dataset(session, "v1")
        assert stats["meetings"] == 1
        assert stats["tdocs"] == 2
        assert stats["nodes"] >= 8
        assert stats["edges"] >= 8
        await activate_dataset(session, "v1")
        await session.commit()
    async with sessions() as session:
        assert await session.scalar(select(func.count(TDocRow.row_id))) == 2
        assert await session.scalar(select(func.count(DocumentBlockRow.row_id))) == 2
        assert await session.scalar(select(func.count(RetrievalChunkRow.row_id))) >= 1
        assert await session.scalar(select(func.count(EvidenceRow.id))) == 4
        assert await session.scalar(select(func.count(KnowledgeNodeRow.row_id))) == stats["nodes"]
        assert await session.scalar(select(func.count(KnowledgeEdgeRow.id))) == stats["edges"]
        dataset = await session.get(DatasetVersionRow, "v1")
        assert dataset and dataset.state == DatasetState.ACTIVE
        artifact_row = await session.scalar(
            select(ArtifactVersionRow).where(ArtifactVersionRow.kind == "tdoc")
        )
        assert artifact_row and artifact_row.parse_status == "parsed"
        assert artifact_row.parse_error is None
    repository = SqlRepository(sessions)
    found, _ = await repository.search_tdocs(SearchRequest(query="carrier"))
    assert [item.id for item in found] == ["R2-2600001"]
    meetings, meeting_cursor = await repository.list_meetings(
        ["RAN2"],
        SearchRequest(filters=SearchFilters(temporal=TemporalScope(last_k_meetings=1)), top_k=1),
    )
    assert [item.id for item in meetings] == ["RAN2-133"]
    assert meeting_cursor is None
    selected = await repository.get_tdoc("R2-2600001")
    assert selected and selected.status.value == "agreed"
    assert selected.specifications == ["38.306"]
    evidence = await repository.evidence(selected.evidence_ids)
    assert evidence and evidence[0].tdoc_id == "R2-2600001"
    blocks = await repository.document_blocks("R2-2600001")
    assert blocks[0].section_path == ["Carrier aggregation"]
    passages = await repository.search_passages(
        "proposal agreed", tdoc_ids=["R2-2600001"], meeting_ids=[], top_k=5
    )
    assert passages and passages[0].evidence_ids

    service = KnowledgeService(repository)
    detail = await service.get_tdoc_detail("R2-2600001")
    assert detail.data and len(detail.data.blocks) == 2
    first_page = await service.get_tdoc_detail("R2-2600001", block_limit=1)
    assert first_page.data and len(first_page.data.blocks) == 1
    assert first_page.completeness == "partial"
    assert first_page.next_cursor
    second_page = await service.get_tdoc_detail(
        "R2-2600001", block_limit=1, cursor=first_page.next_cursor
    )
    assert second_page.data and len(second_page.data.blocks) == 1
    assert second_page.completeness == "complete"
    assert second_page.next_cursor is None
    section_tree = await service.document_section_tree("R2-2600001")
    assert [node["title"] for node in section_tree.data] == ["Document", "Carrier aggregation"]
    revision = await service.revision_chain("R2-2600001")
    assert revision.data == ["R2-2600001"]
    packet = await service.newsletter_packet("RAN2-133", "final")
    assert packet.data and packet.data.totals["tdocs"] == 2
    assert packet.completeness == "complete"
    await engine.dispose()
