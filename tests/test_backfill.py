import hashlib
import io
import zipfile
from datetime import date

import pytest
from docx import Document
from openpyxl import Workbook
from sqlalchemy import select
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from threegpp_kg import backfill
from threegpp_kg.backfill import (
    BackfillError,
    BackfillRequest,
    _directory_names,
    _latest_completed_meeting_ids,
    _required_download,
    _select_report,
    _text,
    _validate_activation_readiness,
    _workbook_candidate_key,
    parse_meeting_date_range,
    run_backfill,
)
from threegpp_kg.config import load_settings, load_working_groups
from threegpp_kg.constants import ArtifactKind, SourceRole
from threegpp_kg.ingestion.download import DownloadedArtifact
from threegpp_kg.sources.adapter import DiscoveredArtifact, DiscoveredMeeting, SourceAdapter
from threegpp_kg.storage.database import ArtifactVersionRow, Base, MeetingRow
from threegpp_kg.storage.object_store import LocalObjectStore


def artifact(filename: str) -> DiscoveredArtifact:
    return DiscoveredArtifact(
        kind=ArtifactKind.REPORT,
        url=f"https://www.3gpp.org/{filename}",
        filename=filename,
        meeting_id="RAN2-134",
    )


def test_meeting_date_range_parses_invitation_and_compact_forms() -> None:
    invitation = "Meetings from Monday 18 to Friday 22 of May 2026 in Dalian, China"
    assert parse_meeting_date_range(invitation) == (
        date(2026, 5, 18),
        date(2026, 5, 22),
    )
    assert parse_meeting_date_range("RAN2 meeting 14-18 October 2024") == (
        date(2024, 10, 14),
        date(2024, 10, 18),
    )
    assert parse_meeting_date_range("14 - 18 October, 2024") == (
        date(2024, 10, 14),
        date(2024, 10, 18),
    )
    assert parse_meeting_date_range("Goa, India, 09/02/2026 to 13/02/2026") == (
        date(2026, 2, 9),
        date(2026, 2, 13),
    )
    assert parse_meeting_date_range("E-meeting, 24 June - 01 July 2026") == (
        date(2026, 6, 24),
        date(2026, 7, 1),
    )
    assert parse_meeting_date_range("R2-260xxxx Maastricht, Aug. 24th - 28th") == (
        date(2026, 8, 24),
        date(2026, 8, 28),
    )
    assert parse_meeting_date_range("No dates available") == (None, None)


def test_report_selection_prefers_latest_substantive_revision() -> None:
    selected = _select_report(
        [
            artifact("R2_134_Skeleton_Report_v20.zip"),
            artifact("Draft_RAN2_134_Meeting_Report_v9.zip"),
            artifact("Draft_RAN2_134_Meeting_Report_v10.zip"),
        ]
    )
    assert selected is not None
    assert selected.filename == "Draft_RAN2_134_Meeting_Report_v10.zip"


def test_report_selection_falls_back_to_skeleton_and_handles_empty() -> None:
    skeleton = artifact("Skeleton_Report_v2.zip")
    assert _select_report([skeleton]) == skeleton
    assert _select_report([]) is None


def test_workbook_selection_prefers_direct_xlsx_then_packaged_xlsx() -> None:
    html = artifact("TdocsByAgenda.htm")
    package = artifact("SA2-173_Index_2026.zip")
    workbook = artifact("TDoc_List_SA2-173.xlsx")
    assert max([html, package], key=_workbook_candidate_key) == package
    assert max([html, package, workbook], key=_workbook_candidate_key) == workbook


@pytest.mark.asyncio
async def test_latest_completed_meetings_exclude_future_placeholders(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    meetings = [
        DiscoveredMeeting(
            id=f"RAN2-{number}",
            working_group_id="RAN2",
            number=number,
            variant="",
            url=f"https://www.3gpp.org/meeting/{number}/",
            source_name=f"TSGR2_{number}",
        )
        for number in range(1, 8)
    ]

    async def discover(*args: object) -> dict[str, list[DiscoveredArtifact]]:
        meeting = next(item for item in meetings if item == args[-1])
        if meeting.number > 5:
            return {"documents": [], "tdoc_lists": [], "reports": []}
        return {
            "documents": [
                DiscoveredArtifact(
                    ArtifactKind.TDOC_LIST,
                    f"{meeting.url}TDoc_List.xlsx",
                    "TDoc_List.xlsx",
                    meeting.id,
                ),
                DiscoveredArtifact(
                    ArtifactKind.TDOC,
                    f"{meeting.url}R2-2600001.zip",
                    "R2-2600001.zip",
                    meeting.id,
                ),
            ],
            "reports": [
                DiscoveredArtifact(
                    ArtifactKind.REPORT,
                    f"{meeting.url}report.zip",
                    "report.zip",
                    meeting.id,
                )
            ],
        }

    monkeypatch.setattr(backfill, "_discover_meeting_artifacts", discover)
    selected = await _latest_completed_meeting_ids(
        load_working_groups()["RAN2"],
        object(),  # type: ignore[arg-type]
        object(),  # type: ignore[arg-type]
        meetings,
        3,
    )

    assert selected == ["RAN2-3", "RAN2-4", "RAN2-5"]


def test_backfill_helpers_fail_closed_and_validate_activation() -> None:
    assert _directory_names('<a href="Docs/">Docs/</a><a href="Reports/">Reports</a>') == {
        "docs",
        "reports",
    }
    raw = downloaded("https://www.3gpp.org/list", b"text\xff", "text/html")
    assert _text(raw) == "text�"
    assert parse_meeting_date_range("32 - 40 Invalid, 2026") == (None, None)

    ready = {
        "meeting_id": "RAN2-134",
        "row_persistence_ratio": 1.0,
        "actionable_archive_match_ratio": 1.0,
        "report": {"filename": "report.docx"},
        "document_results": [],
    }
    _validate_activation_readiness([ready])
    failed = {
        **ready,
        "row_persistence_ratio": 0.9,
        "actionable_archive_match_ratio": 0.9,
        "report": {"error": "bad report"},
        "document_results": [{"tdoc_id": "R2-1", "status": "failed"}],
    }
    with pytest.raises(BackfillError, match="persisted fewer.*actionable.*report.*R2-1"):
        _validate_activation_readiness([failed])


@pytest.mark.asyncio
async def test_required_download_and_backfill_configuration_guards() -> None:
    item = downloaded("https://www.3gpp.org/item", b"ok", "text/plain")

    class Downloader:
        def __init__(self, value: DownloadedArtifact | None) -> None:
            self.value = value

        async def download(self, _url: str) -> DownloadedArtifact | None:
            return self.value

    assert await _required_download(Downloader(item), item.url) == item  # type: ignore[arg-type]
    with pytest.raises(BackfillError, match="not-modified"):
        await _required_download(Downloader(None), item.url)  # type: ignore[arg-type]

    settings = load_settings()
    group = load_working_groups()["RAN2"]
    with pytest.raises(BackfillError, match="does not match"):
        await run_backfill(settings, group, BackfillRequest("RAN3", [], "fixture"))
    with pytest.raises(BackfillError, match="document_limit"):
        await run_backfill(
            settings,
            group,
            BackfillRequest("RAN2", [], "fixture", document_limit=-2),
        )
    s3_settings = settings.model_copy(
        update={
            "object_store": settings.object_store.model_copy(
                update={"backend": "s3", "endpoint": "https://objects.internal"}
            )
        }
    )
    with pytest.raises(BackfillError, match="local object store"):
        await run_backfill(s3_settings, group, BackfillRequest("RAN2", [], "fixture"))
    with pytest.raises(BackfillError, match="PostgreSQL"):
        await run_backfill(settings, group, BackfillRequest("RAN2", [], "fixture"))


def downloaded(url: str, content: bytes, content_type: str) -> DownloadedArtifact:
    return DownloadedArtifact(
        url=url,
        content=content,
        sha256=hashlib.sha256(content).hexdigest(),
        content_type=content_type,
        etag='"fixture"',
        last_modified="Sat, 29 Aug 2026 00:00:00 GMT",
    )


@pytest.mark.asyncio
async def test_nested_inbox_sources_are_discovered_case_insensitively() -> None:
    group = load_working_groups()["RAN2"]
    meeting = DiscoveredMeeting(
        id="RAN2-132",
        working_group_id="RAN2",
        number=132,
        variant="",
        url="https://www.3gpp.org/meeting/",
        source_name="TSGR2_132",
    )
    pages = {
        meeting.url: '<a href="INBOX/">INBOX/</a>',
        f"{meeting.url}INBOX/": (
            '<a href="Chair_Notes/">Chair_Notes/</a>'
            '<a href="Email_Discussions/">Email_Discussions/</a>'
        ),
        f"{meeting.url}INBOX/Chair_Notes/": (
            '<a href="R2_132_ChairNotes_final.docx">R2_132_ChairNotes_final.docx</a>'
        ),
        f"{meeting.url}INBOX/Email_Discussions/": (
            '<a href="RAN2_132_Post_email_discussions_v00.docx">'
            "RAN2_132_Post_email_discussions_v00.docx</a>"
        ),
    }

    class Downloader:
        async def download(self, url: str, **_kwargs: object) -> DownloadedArtifact:
            return downloaded(url, pages[url].encode(), "text/html")

    discovered = await backfill._discover_meeting_artifacts(
        group,
        SourceAdapter(group, {"www.3gpp.org"}),
        Downloader(),  # type: ignore[arg-type]
        meeting,
    )
    assert discovered["chair_notes"][0].kind == ArtifactKind.CHAIR_NOTES
    assert discovered["chair_notes"][0].source_role == SourceRole.CHAIR_NOTES
    assert (
        discovered["post_meeting_discussion"][0].kind
        == ArtifactKind.POST_MEETING_DISCUSSION
    )


def backfill_workbook() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["TDoc", "Title", "Source", "Result"])
    sheet.append(["R2-2600001", "Carrier aggregation", "Qualcomm", "Agreed"])
    sheet.append(["R2-2600002", "Sidelink", "Ericsson", "Agreed"])
    sheet.append(["R2-2600003", "Macro attachment", "Nokia", "Agreed"])
    sheet.append(["R2-2600004", "Unsupported attachment", "Samsung", "Agreed"])
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def report_docx() -> bytes:
    document = Document()
    document.add_heading("Report of RAN2 meeting #134", level=1)
    document.add_paragraph("18 - 22 May, 2026")
    document.add_paragraph("Conclusion: The meeting report was approved.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def tdoc_zip() -> bytes:
    document = Document()
    document.add_heading("Proposal", level=1)
    document.add_paragraph("Agreement: Adopt the proposal.")
    inner = io.BytesIO()
    document.save(inner)
    outer = io.BytesIO()
    with zipfile.ZipFile(outer, "w") as archive:
        archive.writestr("proposal.docx", inner.getvalue())
    return outer.getvalue()


def macro_tdoc_zip() -> bytes:
    outer = io.BytesIO()
    with zipfile.ZipFile(outer, "w") as archive:
        archive.writestr("proposal.docm", b"macro-enabled package")
    return outer.getvalue()


def unsupported_tdoc_zip() -> bytes:
    outer = io.BytesIO()
    with zipfile.ZipFile(outer, "w") as archive:
        archive.writestr("readme.bin", b"no supported document")
    return outer.getvalue()


@pytest.mark.asyncio
async def test_meeting_ingestion_handles_report_date_batches_and_document_failure(
    tmp_path, monkeypatch: pytest.MonkeyPatch
) -> None:
    group = load_working_groups()["RAN2"]
    source_meeting = DiscoveredMeeting(
        id="RAN2-134",
        working_group_id="RAN2",
        number=134,
        variant="",
        url="https://www.3gpp.org/meeting/",
        source_name="TSGR2_134",
    )
    workbook_source = DiscoveredArtifact(
        ArtifactKind.TDOC_LIST,
        "https://www.3gpp.org/list.xlsx",
        "TDoc_List.xlsx",
        source_meeting.id,
    )
    report_source = DiscoveredArtifact(
        ArtifactKind.REPORT,
        "https://www.3gpp.org/report.docx",
        "Final_Report.docx",
        source_meeting.id,
    )
    agenda_source = DiscoveredArtifact(
        ArtifactKind.AGENDA,
        "https://www.3gpp.org/agenda.csv",
        "agenda.csv",
        source_meeting.id,
    )
    documents = [
        DiscoveredArtifact(
            ArtifactKind.TDOC,
            f"https://www.3gpp.org/R2-260000{number}.zip",
            f"R2-260000{number}.zip",
            source_meeting.id,
        )
        for number in (1, 2, 3, 4)
    ]
    discovered = {
        "tdoc_lists": [workbook_source],
        "documents": [workbook_source, *documents],
        "reports": [report_source],
        "agenda": [agenda_source],
        "invitations": [],
    }

    async def discover(*_args: object) -> dict[str, list[DiscoveredArtifact]]:
        return discovered

    monkeypatch.setattr(backfill, "_discover_meeting_artifacts", discover)
    artifacts = {
        workbook_source.url: downloaded(
            workbook_source.url,
            backfill_workbook(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ),
        report_source.url: downloaded(
            report_source.url,
            report_docx(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        agenda_source.url: downloaded(agenda_source.url, b"item,title\n1,Opening\n", "text/csv"),
        documents[0].url: downloaded(documents[0].url, tdoc_zip(), "application/zip"),
        documents[2].url: downloaded(documents[2].url, macro_tdoc_zip(), "application/zip"),
        documents[3].url: downloaded(documents[3].url, unsupported_tdoc_zip(), "application/zip"),
    }

    class Downloader:
        async def download(self, url: str, **kwargs: object) -> DownloadedArtifact | None:
            if url == documents[1].url:
                raise RuntimeError("source unavailable")
            if kwargs.get("etag"):
                return None
            return artifacts[url]

    engine = create_async_engine("sqlite+aiosqlite://")
    sessions = async_sessionmaker(engine, expire_on_commit=False)
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)
    settings = load_settings()
    settings = settings.model_copy(
        update={"parsers": settings.parsers.model_copy(update={"document_workers": 1})}
    )
    result = await backfill._ingest_meeting(
        settings,
        group,
        SourceAdapter(group, {"www.3gpp.org"}),
        Downloader(),  # type: ignore[arg-type]
        sessions,
        LocalObjectStore(tmp_path / "objects"),
        source_meeting,
        BackfillRequest("RAN2", ["RAN2-134"], "fixture", document_limit=-1),
    )
    assert result["starts_on"] == "2026-05-18"
    assert result["ends_on"] == "2026-05-22"
    assert result["report"]["date_source"] == "report"
    assert [item["status"] for item in result["document_results"]] == [
        "ingested",
        "failed",
        "quarantined",
        "failed",
    ]
    async with sessions() as session:
        meeting = await session.scalar(select(MeetingRow))
        assert meeting and meeting.ends_on == date(2026, 5, 22)
        quarantined = await session.scalar(
            select(ArtifactVersionRow).where(ArtifactVersionRow.parse_status == "quarantined")
        )
        assert quarantined and quarantined.filename == "R2-2600003.zip"
        failed = await session.scalar(
            select(ArtifactVersionRow).where(ArtifactVersionRow.parse_status == "failed")
        )
        assert failed and failed.filename == "R2-2600004.zip"

    repeated = await backfill._ingest_meeting(
        settings,
        group,
        SourceAdapter(group, {"www.3gpp.org"}),
        Downloader(),  # type: ignore[arg-type]
        sessions,
        LocalObjectStore(tmp_path / "objects"),
        source_meeting,
        BackfillRequest("RAN2", ["RAN2-134"], "fixture", document_limit=-1),
    )
    assert [item["status"] for item in repeated["document_results"]] == [
        "unchanged",
        "failed",
        "quarantined",
        "failed",
    ]
    await engine.dispose()
